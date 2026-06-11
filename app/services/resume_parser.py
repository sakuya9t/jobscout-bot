"""Extract normalized plain text from an uploaded resume (PDF / docx / txt / md)."""
from __future__ import annotations

import io
import re


def _clean(text: str) -> str:
    # Collapse runs of whitespace/newlines so prompts stay compact.
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _from_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on extension. Raises ValueError for unsupported types, empty
    text, or an unreadable (corrupt/encrypted) file — so the upload route can
    return 422 instead of a 500. pypdf (PdfReadError) and python-docx
    (BadZipFile/PackageNotFoundError) raise their own exception types on
    malformed input; normalize them all to ValueError here."""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            text = _from_pdf(data)
        elif name.endswith(".docx"):
            text = _from_docx(data)
        elif name.endswith((".txt", ".md", ".markdown")):
            text = data.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported resume type: {filename!r} (use pdf, docx, txt, md)")
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001 — any parser failure means a bad file
        raise ValueError(
            f"Could not read {filename!r}; the file may be corrupt or password-protected."
        ) from exc

    text = _clean(text)
    if not text:
        raise ValueError("Could not extract any text from the resume.")
    return text
